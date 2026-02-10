import docx
import pandas as pd

from docx.shared import Inches, Cm, Mm

def new_word_file(last_trading_day):
    """Creates a new word file named Daily Marke Report As on {last_trading_day}"""

    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor

    doc = docx.Document()
    paragraph = doc.add_paragraph(f'------------------------------\nDaily Market Report \nAs on \n{last_trading_day}\n------------------------------')
    paragraph_format = paragraph.paragraph_format

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    # font.bold = True

    font.color.rgb = RGBColor(7, 2, 56)

    paragraph.style = doc.styles['Normal']

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(300)

    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    paragraph_format.space_before, paragraph_format.space_after
    paragraph_format.space_before = Pt(160)
    paragraph_format.alignment
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.save(f'/Users/arun/Documents/Python Projects/Daily Momentum Report/Report as on {last_trading_day}.docx')
#========================================================================================================================================================
#========================================================================================================================================================

def df_to_word(last_trading_day, table_title, df):
    """This function writes dataframe to word. last_trading_day is needed to save the doc file."""

    pd.options.display.float_format = '{:.2f}'.format
    # open an existing document
#     doc = Document()

    doc = docx.Document(f'/Users/arun/Documents/Python Projects/Daily Momentum Report/Report as on {last_trading_day}.docx')
    
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row

    doc.add_paragraph(table_title)

    t = doc.add_table(df.shape[0]+1, df.shape[1], style = 'Light Shading Accent 1')

    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    doc.add_paragraph('')
    doc.add_paragraph('')
    # save the doc
    doc.save(f'/Users/arun/Documents/Python Projects/Daily Momentum Report/Report as on {last_trading_day}.docx')

#========================================================================================================================================================
#========================================================================================================================================================
def chart_to_word(last_trading_day, chart_title, chart_file_name, ltp,volume, percent_turn, sup, rest,  remarks = False, big = False):
    """ Inserts chart to word from D:/Charts"""

    doc = docx.Document(f'/Users/arun/Documents/Python Projects/Daily Momentum Report/Report as on {last_trading_day}.docx')
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_text(chart_title)
    if big == True:
        r.add_picture(f'/Users/arun/Documents/Python Projects/Daily Momentum Report/Charts/{chart_file_name}')#, width=Inches(10), height = Inches(9.2))
    else:
        r.add_picture(f'/Users/arun/Documents/Python Projects/Daily Momentum Report/Charts/{chart_file_name}')#, width=Inches(10), height = Inches(6))#
        
    if remarks == True:
        r.add_text('Remarks:')
        r.add_break()
        r.add_text(f'Current Level : {ltp}')
        r.add_break()
        volume= volume/10000000
        r.add_text(f'Turnover : {volume.round(decimals=3)} Crores [{percent_turn}%]')
        r.add_break()
        r.add_text(f'Support : {sup}')
        r.add_break()
        r.add_text(f'Resistance : {rest}')
    else:
        pass
#     doc.add_page_break()
    doc.save(f'/Users/arun/Documents/Python Projects/Daily Momentum Report/Report as on {last_trading_day}.docx')
# ======================================================================================================================================== 
def new_word_file_for_intradayreport():
    """Creates a new word file named Daily Marke Report As on {last_trading_day}"""

    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
    import datetime

    current_time = datetime.datetime.today()
    rounded_time = current_time.strftime("%Y-%m-%d %H:%M:%S")
    doc = docx.Document()
    paragraph = doc.add_paragraph(f'------------------------------\nIntraday Report \nAs at \n{rounded_time}\n----------------------------------')
    paragraph_format = paragraph.paragraph_format

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    # font.bold = True

    font.color.rgb = RGBColor(7, 2, 56)

    paragraph.style = doc.styles['Normal']

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(300)

    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    paragraph_format.space_before, paragraph_format.space_after
    paragraph_format.space_before = Pt(160)
    paragraph_format.alignment
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.save(f'C:\\Users\\Dell\\Reports\\Live Intraday Report.docx')
    
    
def chart_to_word_forintraday(last_trading_day, chart_title, chart_file_name, ltp,volume, percent_turn, sup, rest,  remarks = False, big = False):
    """ Inserts chart to word from D:/Charts"""

    doc = docx.Document(f'C:\\Users\\Dell\\Reports\\Live Intraday Report.docx')
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_text(chart_title)
    if big == True:
        r.add_picture(f'D:\\Charts\\{chart_file_name}', width=Inches(10), height = Inches(9.2))#
    else:
        r.add_picture(f'D:\\Charts\\{chart_file_name}', width=Inches(10), height = Inches(6))#
        
    if remarks == True:
        r.add_text('Remarks:')
        r.add_break()
        r.add_text(f'Current Level : {ltp}')
        r.add_break()
        volume= volume/10000000
        r.add_text(f'Turnover : {volume.round(decimals=3)} Crores [{percent_turn}%]')
        r.add_break()
        r.add_text(f'Support : {sup}')
        r.add_break()
        r.add_text(f'Resistance : {rest}')
    else:
        pass
#     doc.add_page_break()
    doc.save(f'C:\\Users\\Dell\\Reports\\Live Intraday Report.docx')
